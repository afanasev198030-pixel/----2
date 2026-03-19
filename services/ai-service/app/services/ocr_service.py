"""
OCR service — document text extraction.

Routing:
  PDF (text-based) → enhanced pdfplumber (find_tables + outside_bbox + extract_text).
  PDF (scanned)    → Vision OCR (DeepSeek-OCR-2) → legacy Tesseract fallback.
  Images           → Vision OCR → legacy Tesseract fallback.
  Excel            → openpyxl / xlrd.

Rollback: set OCR_ENABLED=false in .env → disables Vision OCR for scans.
"""
import io
import time as _time

import pdfplumber
import structlog

from app.utils.text_processing import clean_ocr_text

logger = structlog.get_logger()

try:
    import pytesseract
    from PIL import Image
    HAS_TESSERACT = True
except ImportError:
    HAS_TESSERACT = False
    logger.info("pytesseract_not_available", msg="OCR for scanned documents disabled")


# ---------------------------------------------------------------------------
# Main entry points
# ---------------------------------------------------------------------------

def extract_text(file_bytes: bytes, filename: str) -> str:
    """Extract text from document.

    Routing:
    - Excel -> openpyxl/xlrd (unchanged)
    - PDF   -> enhanced pdfplumber first (tables + text); if scanned -> Vision OCR -> legacy
    - Image -> Vision OCR -> legacy (Tesseract)
    """
    from app.config import get_settings
    settings = get_settings()
    filename_lower = filename.lower()

    if filename_lower.endswith(('.xlsx', '.xls')):
        return extract_text_from_excel(file_bytes, filename)

    if filename_lower.endswith(('.docx', '.doc')):
        return _extract_text_from_docx(file_bytes, filename)

    if filename_lower.endswith('.pdf'):
        enhanced = _extract_pdf_enhanced(file_bytes, filename)
        if enhanced:
            return enhanced
        if settings.has_vision_ocr:
            try:
                return _extract_with_vision_ocr(file_bytes, filename)
            except Exception as e:
                logger.warning(
                    "vision_ocr_failed_fallback",
                    error=str(e)[:200],
                    filename=filename,
                    msg="Scanned PDF: Vision OCR failed, falling back to legacy",
                )
        return _extract_text_legacy(file_bytes, filename)

    if settings.has_vision_ocr:
        try:
            return _extract_with_vision_ocr(file_bytes, filename)
        except Exception as e:
            logger.warning(
                "vision_ocr_failed_fallback",
                error=str(e)[:200],
                filename=filename,
                msg="Image: Vision OCR failed, falling back to legacy",
            )
    return _extract_text_legacy(file_bytes, filename)


def extract_text_debug(file_bytes: bytes, filename: str) -> dict:
    """Extract text with debug metadata — enhanced, vision, and legacy results."""
    from app.config import get_settings
    settings = get_settings()
    filename_lower = filename.lower()
    t0 = _time.monotonic()

    enhanced_result = None
    vision_result = None
    legacy_result = None
    chosen_method = "unknown"
    chosen_reason = ""

    is_docx = filename_lower.endswith(('.docx', '.doc'))
    is_pdf = filename_lower.endswith('.pdf')

    if is_docx:
        td0 = _time.monotonic()
        docx_text = _extract_text_from_docx(file_bytes, filename)
        return {
            "method": "docx_parser",
            "reason": "DOCX document — direct text extraction",
            "text": docx_text,
            "chars": len(docx_text),
            "pages": 1,
            "duration_ms": int((_time.monotonic() - td0) * 1000),
            "enhanced": None,
            "vision": None,
            "legacy": None,
        }

    if is_pdf:
        te0 = _time.monotonic()
        enhanced_text = _extract_pdf_enhanced(file_bytes, filename)
        enhanced_result = {
            "text": enhanced_text,
            "method": "enhanced_pdfplumber",
            "chars": len(enhanced_text),
            "duration_ms": int((_time.monotonic() - te0) * 1000),
        }

    if settings.has_vision_ocr and not filename_lower.endswith(('.xlsx', '.xls')):
        try:
            tv0 = _time.monotonic()
            vision_text = _extract_with_vision_ocr(file_bytes, filename)
            vision_result = {
                "text": vision_text,
                "method": "vision_ocr",
                "chars": len(vision_text),
                "duration_ms": int((_time.monotonic() - tv0) * 1000),
            }
        except Exception as e:
            vision_result = {"error": str(e)[:200], "method": "vision_ocr"}

    tl0 = _time.monotonic()
    legacy_text, legacy_method, legacy_pages = _extract_text_legacy_debug(file_bytes, filename)
    legacy_result = {
        "text": legacy_text,
        "method": legacy_method,
        "pages": legacy_pages,
        "chars": len(legacy_text),
        "duration_ms": int((_time.monotonic() - tl0) * 1000),
    }

    if is_pdf and enhanced_result and enhanced_result["chars"] > 0:
        primary = enhanced_result
        chosen_method = "enhanced_pdfplumber"
        chosen_reason = "Text-based PDF: tables via find_tables(), text via extract_text()"
    elif vision_result and "error" not in vision_result:
        primary = vision_result
        chosen_method = "vision_ocr"
        chosen_reason = "Scanned PDF or image: Vision LLM extraction"
    else:
        primary = legacy_result
        chosen_method = legacy_method
        chosen_reason = "Fallback: legacy pdfplumber/Tesseract"

    duration_ms = int((_time.monotonic() - t0) * 1000)
    return {
        "text": primary["text"],
        "method": primary["method"],
        "chosen_method": chosen_method,
        "chosen_reason": chosen_reason,
        "pages": legacy_result.get("pages", 0),
        "chars": primary["chars"],
        "duration_ms": duration_ms,
        "ocr_enhanced": enhanced_result,
        "ocr_vision": vision_result,
        "ocr_legacy": legacy_result,
    }


# ---------------------------------------------------------------------------
# Vision OCR (DeepSeek-OCR-2)
# ---------------------------------------------------------------------------

def _extract_with_vision_ocr(file_bytes: bytes, filename: str) -> str:
    """Route file through DeepSeek-OCR-2 Vision LLM."""
    from app.services.ocr_client import ocr_pdf_to_markdown, ocr_image_to_markdown

    filename_lower = filename.lower()

    if filename_lower.endswith('.pdf'):
        return ocr_pdf_to_markdown(file_bytes, filename)
    elif filename_lower.endswith(('.jpg', '.jpeg', '.png', '.tiff', '.tif', '.bmp')):
        return ocr_image_to_markdown(file_bytes, filename)
    else:
        return ocr_pdf_to_markdown(file_bytes, filename)


# ---------------------------------------------------------------------------
# Enhanced pdfplumber (text-based PDF) — tables + text without duplication
# ---------------------------------------------------------------------------

_MIN_CHARS_PER_PAGE = 50


def _table_to_markdown(table: list[list]) -> str:
    """Convert pdfplumber table (list of rows with possible None cells) to Markdown."""
    if not table or not table[0]:
        return ""
    rows: list[list[str]] = []
    for row in table:
        cells = [(str(c) if c else "").strip().replace("|", "/") for c in row]
        if any(cells):
            rows.append(cells)
    if not rows:
        return ""
    max_cols = max(len(r) for r in rows)
    for r in rows:
        while len(r) < max_cols:
            r.append("")
    header = "| " + " | ".join(rows[0]) + " |"
    sep = "| " + " | ".join("---" for _ in rows[0]) + " |"
    body = "\n".join("| " + " | ".join(r) + " |" for r in rows[1:])
    return header + "\n" + sep + "\n" + body


def _extract_pdf_enhanced(file_bytes: bytes, filename: str) -> str:
    """Extract text from text-based PDF using pdfplumber with table support.

    Strategy:
    - dedupe_chars() removes overlapping glyphs (bold rendered as double chars).
    - Always extract full page text via extract_text(layout=True) — preserves
      column alignment and never clips header/footer content.
    - find_tables() locates tables; if found, append structured Markdown
      as supplementary data (LLM handles minor duplication gracefully).
    - Returns empty string when avg chars/page < threshold (scanned PDF).
    """
    try:
        parts: list[str] = []
        total_pages = 0
        tables_found = 0

        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            if not pdf.pages:
                return ""
            total_pages = len(pdf.pages)

            for page in pdf.pages:
                page_parts: list[str] = []
                page_to_read = page.dedupe_chars()

                text = page_to_read.extract_text(layout=True)
                if text and text.strip():
                    page_parts.append(text.strip())

                tables = page_to_read.find_tables()
                if tables:
                    tables_found += len(tables)
                    for table in tables:
                        data = table.extract()
                        md = _table_to_markdown(data)
                        if md:
                            page_parts.append(md)

                if page_parts:
                    parts.append("\n\n".join(page_parts))

        combined = "\n\n---\n\n".join(parts)
        avg_chars = len(combined.strip()) / max(total_pages, 1)

        if avg_chars >= _MIN_CHARS_PER_PAGE:
            logger.info(
                "ocr_enhanced_pdfplumber",
                filename=filename,
                pages=total_pages,
                tables_found=tables_found,
                chars=len(combined),
                avg_chars_per_page=int(avg_chars),
            )
            return clean_ocr_text(combined)

        logger.info(
            "enhanced_pdfplumber_insufficient",
            filename=filename,
            pages=total_pages,
            avg_chars_per_page=int(avg_chars),
            threshold=_MIN_CHARS_PER_PAGE,
            msg="PDF appears scanned, deferring to Vision OCR",
        )
        return ""
    except Exception as e:
        logger.warning(
            "enhanced_pdfplumber_failed",
            filename=filename,
            error=str(e)[:200],
        )
        return ""


# ---------------------------------------------------------------------------
# Legacy OCR (pdfplumber + Tesseract) — fully preserved for fallback
# ---------------------------------------------------------------------------

def _extract_text_legacy(file_bytes: bytes, filename: str) -> str:
    """Legacy dispatcher — pdfplumber/Tesseract/openpyxl."""
    filename_lower = filename.lower()
    if filename_lower.endswith('.pdf'):
        return _extract_text_from_pdf_legacy(file_bytes)
    elif filename_lower.endswith(('.jpg', '.jpeg', '.png', '.tiff', '.bmp')):
        return _extract_text_from_image_legacy(file_bytes)
    elif filename_lower.endswith(('.xlsx', '.xls')):
        return extract_text_from_excel(file_bytes, filename)
    else:
        logger.warning("unknown_file_type_legacy", filename=filename)
        return _extract_text_from_pdf_legacy(file_bytes)


def _extract_text_from_pdf_legacy(file_bytes: bytes) -> str:
    """Extract text from PDF using pdfplumber. Fallback to Tesseract OCR."""
    try:
        text_parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

        text = "\n".join(text_parts)

        if (not text or len(text.strip()) < 10) and HAS_TESSERACT:
            logger.info("pdf_no_text_trying_ocr")
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    try:
                        img = page.to_image(resolution=200)
                        if img:
                            ocr_text = pytesseract.image_to_string(img.original, lang='rus+eng')
                            if ocr_text:
                                text_parts.append(ocr_text)
                    except Exception as e:
                        logger.warning("ocr_page_failed", error=str(e))
            text = "\n".join(text_parts)

        return clean_ocr_text(text)
    except Exception as e:
        logger.error("pdf_extraction_failed", error=str(e))
        return ""


def _extract_text_from_image_legacy(file_bytes: bytes) -> str:
    """Extract text from image using pytesseract."""
    if not HAS_TESSERACT:
        logger.warning("tesseract_not_available_for_image")
        return ""
    try:
        image = Image.open(io.BytesIO(file_bytes))
        text = pytesseract.image_to_string(image, lang='rus+eng')
        return clean_ocr_text(text)
    except Exception as e:
        logger.error("image_extraction_failed", error=str(e))
        return ""


def _extract_text_legacy_debug(file_bytes: bytes, filename: str) -> tuple[str, str, int]:
    """Legacy extraction with debug info. Returns (text, method, pages)."""
    filename_lower = filename.lower()
    pages = 0

    if filename_lower.endswith('.pdf'):
        method = "pdfplumber"
        try:
            text_parts = []
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                pages = len(pdf.pages)
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            text = "\n".join(text_parts)
            if (not text or len(text.strip()) < 10) and HAS_TESSERACT:
                method = "tesseract"
                text_parts = []
                with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                    for page in pdf.pages:
                        try:
                            img = page.to_image(resolution=200)
                            if img:
                                ocr_text = pytesseract.image_to_string(img.original, lang='rus+eng')
                                if ocr_text:
                                    text_parts.append(ocr_text)
                        except Exception:
                            pass
                text = "\n".join(text_parts)
            text = clean_ocr_text(text)
        except Exception as e:
            text = ""
            method = f"pdfplumber_error: {str(e)[:100]}"
    elif filename_lower.endswith(('.jpg', '.jpeg', '.png', '.tiff', '.bmp')):
        method = "tesseract"
        pages = 1
        text = _extract_text_from_image_legacy(file_bytes)
    elif filename_lower.endswith(('.xlsx', '.xls')):
        method = "openpyxl" if filename_lower.endswith('.xlsx') else "xlrd"
        text = extract_text_from_excel(file_bytes, filename)
    else:
        method = "pdfplumber_fallback"
        text = _extract_text_from_pdf_legacy(file_bytes)

    return text, method, pages


# ---------------------------------------------------------------------------
# DOCX extraction
# ---------------------------------------------------------------------------

def _extract_text_from_docx(file_bytes: bytes, filename: str) -> str:
    """Extract text from DOCX, deduplicating repeated table cells."""
    import io
    try:
        from docx import Document
    except ImportError:
        logger.warning("python_docx_not_installed", msg="python-docx unavailable, returning empty")
        return ""

    try:
        doc = Document(io.BytesIO(file_bytes))
        seen: set[str] = set()
        lines: list[str] = []
        for p in doc.paragraphs:
            t = p.text.strip()
            if t and t not in seen:
                seen.add(t)
                lines.append(t)
        for table in doc.tables:
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    t = cell.text.strip()
                    if t and t not in seen:
                        seen.add(t)
                        cells.append(t)
                if cells:
                    lines.append(" | ".join(cells))
        text = "\n".join(lines)
        logger.info("docx_extracted", filename=filename, chars=len(text))
        return text
    except Exception as e:
        logger.error("docx_extraction_failed", filename=filename, error=str(e)[:200])
        return ""


# ---------------------------------------------------------------------------
# Excel extraction (shared — not affected by Vision OCR)
# ---------------------------------------------------------------------------

def extract_text_from_excel(file_bytes: bytes, filename: str) -> str:
    """Extract text from Excel (.xlsx/.xls) as Markdown tables."""
    filename_lower = filename.lower()

    if filename_lower.endswith('.xls') and not filename_lower.endswith('.xlsx'):
        return _extract_text_xls(file_bytes, filename)

    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True, read_only=True)
        text_parts = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            text_parts.append(f"## Sheet: {sheet_name}\n")

            rows: list[list[str]] = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(cell).strip() if cell is not None else "" for cell in row]
                if any(c for c in cells):
                    rows.append(cells)

            if rows:
                text_parts.append(_rows_to_markdown_table(rows))

        wb.close()
        text = "\n\n".join(text_parts)
        logger.info("excel_extracted", filename=filename, chars=len(text),
                     sheets=len(wb.sheetnames), format="markdown")
        return text
    except Exception as e:
        logger.error("xlsx_extraction_failed", filename=filename, error=str(e))
        return _extract_text_xls(file_bytes, filename)


def _rows_to_markdown_table(rows: list[list[str]]) -> str:
    """Convert a list of rows into a Markdown table. First row = header."""
    if not rows:
        return ""

    max_cols = max(len(r) for r in rows)
    for r in rows:
        while len(r) < max_cols:
            r.append("")

    header = rows[0]
    header_line = "| " + " | ".join(header) + " |"
    sep_line = "| " + " | ".join("---" for _ in header) + " |"
    body_lines = []
    for row in rows[1:]:
        body_lines.append("| " + " | ".join(row) + " |")

    return "\n".join([header_line, sep_line] + body_lines)


def _extract_text_xls(file_bytes: bytes, filename: str) -> str:
    """Extract text from old .xls format using xlrd, output as Markdown."""
    try:
        import xlrd
        wb = xlrd.open_workbook(file_contents=file_bytes)
        text_parts = []
        for sheet in wb.sheets():
            text_parts.append(f"## Sheet: {sheet.name}\n")
            rows: list[list[str]] = []
            for row_idx in range(sheet.nrows):
                cells = []
                for col_idx in range(sheet.ncols):
                    val = sheet.cell(row_idx, col_idx).value
                    cells.append(str(val).strip() if val is not None else "")
                if any(c for c in cells):
                    rows.append(cells)
            if rows:
                text_parts.append(_rows_to_markdown_table(rows))
        text = "\n\n".join(text_parts)
        logger.info("xls_extracted", filename=filename, chars=len(text),
                     sheets=wb.nsheets, format="markdown")
        return text
    except ImportError:
        logger.error("xlrd_not_installed", filename=filename, msg="pip install xlrd")
        return ""
    except Exception as e:
        logger.error("xls_extraction_failed", filename=filename, error=str(e))
        return ""
